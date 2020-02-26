# coding: utf-8

"""
tojson.py

Usage:
  tojson.py <word_file> [<json_file>]
  tojson.py -h | --help
  tojson.py --version

Options:
  -h --help              Show this screen.
  --version              Show version.
"""

import logging
import sys
import os
import re
import json
import traceback
from datetime import datetime
from typing import List, Dict

from docopt import docopt
from docx import table as dtable
import docx
import pandas as pd

logger = logging.getLogger(__name__)


class ReportInfo(object):
    """报告信息"""
    report_number = None  # 报告编号
    report_query_time = None  # 查询请求时间
    report_create_time = None  # 报告时间
    name = None  # 被查询者姓名
    cert_type = None  # 被查询者证件类型
    cert_no = None  # 被查询者证件号码
    query_operator = None  # 查询操作员
    query_reason = None  # 查询原因


class MessageHeader(object):
    queryTime = None
    reportCreateTime = None
    reportSN = None


class QueryRequest(object):
    name = None
    certtype = None
    certno = None
    userCode = None
    queryReason = None
    # queryResultCue = None
    # queryOrg = None
    # productType = None
    # formatVersion = None
    # format = None


class Header(object):
    messageHeader = None
    queryReq = None


class Identity(object):
    """ 身份信息 """
    gender = None
    birthday = None
    maritalState = None
    mobile = None
    officeTelephoneNo = None
    homeTelephoneNo = None
    eduLevel = None
    eduDegree = None
    postAddress = None
    registeredAddress = None


class Residence(object):
    """ 居住信息"""
    getTime = None
    residenceType = None
    address = None


class Spouse(object):
    """配偶信息"""
    certType = None
    certNo = None
    employer = None
    telephoneNo = None
    name = None


class Professional(object):
    """ 职业信息"""
    getTime = None
    employer = None
    employerAddress = None
    occupation = None
    industry = None
    duty = None
    startYear = None
    title = None


class PersonalBasicInfo(object):
    identity = None
    residence = None
    spouse = None
    professional = None


class CreditCue(object):
    """信用提示"""
    otherLoanCount = None
    firstLoanOpenMonth = None
    loanCardCount = None
    firstLoanCardOpenMonth = None
    standardLoanCardCount = None
    firstStandardLoanCardOpenMonth = None
    announceCount = None
    dissentCount = None
    perHouseLoanCount = None
    perBusinessHouseLoanCount = None


class FellBackSummary(object):
    fellBackDebtSumCount = None
    fellBackDebtSumBalance = None
    assetDispositionSumCount = None
    assetDispositionSumBalance = None
    assureerRepaySumCount = None
    assureerRepaySumBalance = None


class OverdueSummary(object):
    loanSumHighestOverdueAmountPerMon = None
    loanSumMaxDuration = None
    loanSumMonths = None
    loanSumCount = None
    loanCardSumHighestOverdueAmountPerMon = None
    loanCardSumMaxDuration = None
    loanCardSumMonths = None
    loanCardSumCount = None
    standardLoanCardSumHighestOverdueAmountPerMon = None
    standardLoanCardSumMaxDuration = None
    standardLoanCardSumMonths = None
    standardLoanCardSumCount = None


class OverdueAndFellBack(object):
    fellBackSummary = None
    overdueSummary = None


class ShareAndDebtCommon(object):
    balance = None
    latest6MonthUsedAvgAmount = None
    financeCorpCount = None
    financeOrgCount = None
    accountCount = None
    creditLimit = None
    maxCreditLimitPerOrg = None
    minCreditLimitPerOrg = None
    usedCreditLimit = None


class ShareAndDebt(object):
    unPaidLoan = None
    unDestroyLoanCard = None
    unDestroyStandardLoanCard = None


class SummaryInfo(object):
    creditCue = CreditCue()
    overdueAndFellBack = OverdueAndFellBack()
    shareAndDebt = ShareAndDebt()


class OverdueRecordDetail(object):
    lastMonths = None
    amount = None
    month = None


class OverdueRecord(object):
    overdueRecordDetail = None
    dates = None
    # specials = None


class SpecialRecord(object):
    tradeType = None
    date = None
    changeMonths = None
    amount = None
    detail = None


class Loan(object):
    statements = None
    state = None
    class5State = None
    balance = None
    remainPaymentCyc = None
    scheduledPaymentAmount = None
    scheduledPaymentDate = None
    actualPaymentAmount = None
    recentPayDate = None
    currOverdueCyc = None
    currOverdueAmount = None
    overdue31To60Amount = None
    overdue61To90Amount = None
    overdue91To180Amount = None
    overdueOver180Amount = None
    latest24State = None
    latest24Date = None
    settle = False
    overdueRecord = None
    specials = None


class LoanCard(object):
    statements = None
    state = None
    usedCreditLimitAmount = None
    latest6MonthUsedAvgAmount = None
    usedHighestAmount = None
    scheduledPaymentAmount = None
    scheduledPaymentDate = None
    actualPaymentAmount = None
    recentPayDate = None
    currOverdueCyc = None
    currOverdueAmount = None
    latest24State = None
    latest24Date = None
    overdueRecord = None


class AssurerRepay(object):
    org = None
    recentAssurerRepayDate = None
    accumulativeAssurerRepayAmount = None
    recentRepayDate = None
    balance = None


class Guarantee(object):
    organname = None
    contractMoney = None
    beginDate = None
    endDate = None
    guananteeMoney = None
    guaranteeBalance = None
    class5State = None
    billingDate = None


class GuaranteeInfo(object):
    guaranteeFormat = None
    guarantee = None


class CreditDetail(object):
    assurerRepay = None
    loan = None
    loanCard = None
    standardLoanCard = None
    guaranteeInfo = None


class QueryRecordDetail(object):
    queryReason = None
    queryDate = None
    querier = None


class QueryRecordSummary(object):
    latestMonthQueryorgSumLoanApproval = None
    latestMonthQueryorgSumLoanCardApproval = None
    latestMonthQueryRecordSumLoanApproval = None
    latestMonthQueryRecordSumLoanCardApproval = None
    latestMonthQueryRecordSumPersonal = None
    twoYearQueryRecordSumCollection = None
    twoYearQueryRecordSumGuarantee = None
    twoYearQueryRecordSumSpecial = None


class QueryRecord(object):
    recordSummary = None
    recordInfo = None


class AccFund(object):
    area = None
    registerDate = None
    firstMonth = None
    toMonth = None
    state = None
    pay = None
    ownPercent = None
    comPercent = None
    organname = None
    getTime = None


class PublicInfo(object):
    accFund = None


class PBOCEntity(object):
    body_str = None
    header = ReportInfo()
    personalInfo = PersonalBasicInfo()
    creditDetail = CreditDetail()
    queryRecord = None


def obj_to_dict(obj):
    if isinstance(obj, list):
        return [obj_to_dict(o) for o in obj]
    elif not isinstance(obj, dict) and not isinstance(obj, int) and not isinstance(obj, str) and obj is not None:
        obj = obj.__dict__
        for k in obj:
            obj[k] = obj_to_dict(obj[k])
        return obj
    else:
        return obj


def prefix_word(document: docx.Document) -> list:
    """
    将docx解析的报告格式转为list[str|DataFrame]格式,其中描述转为str,数据转为DataFrame
    :param document:
    :return:
    """
    body = read_tables(document)
    return body


def table2dataFrame(table: dtable.Table) -> pd.DataFrame:
    """
    Table转DataFrame
    Table中的cell必须不包含Table
    :param table:
    :return:
    """
    lst = []
    for ii, row in enumerate(table.rows):
        row_lst = []
        for jj, cell in enumerate(row.cells):
            if len(cell.tables) == 0:
                v = re.sub('(^[\n ]*)|([\n ]*$)', '', cell.text)  # 去除特殊字符
                row_lst.append(None if v == '' else v)
            else:
                raise ValueError("cell value contain tables")
        if not check_df_row(row_lst):
            lst.append(row_lst)
    return pd.DataFrame(lst)


def check_df_row(row: list) -> bool:
    """
    判断list的元素是否全为None
    :param row:
    :return:
    """
    for v in row:
        if v is not None:
            return False
    return True


def check_df(df: str, value: str):
    """简单判断两个df是否数据一致"""
    # if not isinstance(value, pd.DataFrame):
    #     return False
    # elif df.shape != value.shape:
    #     return False
    # else:
    #     l = list(value[list(df.columns)[-1]])
    #     r = list(df[list(df.columns)[-1]])
    #     for x, y in zip(l, r):
    #         if x != y:
    #             return False
    # return True
    return df == value


def read_tables(document):
    """
    读取所有的table并做简单处理
    :param document:
    :return:
    """
    body = []
    value = None
    for table in document.tables:
        for ii, cell in enumerate(table._cells):
            if len(cell.tables) == 0:
                if isinstance(value, str) and cell.text.strip() == value:  # 重复内容
                    continue
                body.append(cell.text.strip())
                value = cell.text.strip()
            else:
                assert len(cell.tables) == 1
                df = table2dataFrame(cell.tables[0])  # 转化为DataFrame
                df_str = str(df)
                if check_df(df_str, value):  # 重复table
                    continue
                body.append(df)
                value = df_str
    return body


def read_report_info(body):
    """
    报告基本信息
    --------------------------------
    0,1,2
    个人信用报告,个人信用报告,个人信用报告
    "(银行版) \n","(银行版) \n","(银行版) \n"
    # 报告编号:20180806000059,查询请求时间:2 ,报告时间:2018.9  --remove
    --------------------------------
    0,1,2,3,4
    被查询者姓名 ,被查询者证件类型 ,被查询者证件号码 ,查询操作员 ,查询原因
    ,身份证,,,贷后管理
    --------------------------------
    :param body:
    :return:
    """
    report_info_body = body[:2]
    header = Header()
    message_header = MessageHeader()
    # if len(report_info_body[0].columns) >= 3:
    #     message_header.reportSN = report_info_body[0][0][2].split(':')[-1].strip()
    #     message_header.queryTime = report_info_body[0][1][2].split(':')[-1].strip()
    #     message_header.reportCreateTime = report_info_body[0][2][2].split(':')[-1].strip()
    header.messageHeader = message_header

    query_request_df = report_info_body[1]
    if len(query_request_df) != 0:
        query_request = QueryRequest()
        tags = ['被查询者姓名', '被查询者证件类型', '被查询者证件号码', '查询操作员', '查询原因']
        rs = find_values_from_df_by_group_tags(query_request_df, [tags])[0]
        query_request.name = rs.get('被查询者姓名')
        query_request.certtype = rs.get('被查询者证件类型')
        query_request.certno = rs.get('被查询者证件号码')
        query_request.userCode = rs.get('查询操作员')
        query_request.queryReason = rs.get('查询原因')
        header.queryReq = query_request

    return header


def find_values_from_df_by_group_tags(df: pd.DataFrame, group_tags: List[List[str]]) -> List[dict]:
    """
    根据标签组给的的标签,从DataFrame中查找到标签对应的值并返回
    注意:
    1.DataFrame格式必须为[标题+数据]循环模式
    Example
    --
    >>> df1 = pd.DataFrame([['name', 'age', 'gender'], ['smith', 15, 'male']])
    >>> group_tags1 = [['name', 'age']]
    >>> find_values_from_df_by_group_tags(df1, group_tags1)
    '[{\'name\': \'smith\', \'age\': 15}]'
    :param df:
    :param group_tags:
    :return:

    """
    cls, title = None, None  # 类别及每个类别的标题
    lst = []
    find_title = False
    for ii, row in df.iterrows():
        if not find_title:  # 打标记,当行为标题时,记录标题及类别
            for jj, g in enumerate(group_tags):
                for kk, e in enumerate(row):
                    if pd.isnull(e):
                        continue
                    if g[kk] in e:
                        find_title = True
                        cls, title = jj, row
                        break
        if find_title:
            find_title = False
            continue
        for jj, g in enumerate(group_tags):  # 当行为数据时,记录当前标题对应的数据
            if cls != jj:
                continue
            rs = {}
            for tag in g:
                for kk, v in enumerate(row):
                    if tag in title[kk]:
                        if tag in rs:
                            rs[tag] = [rs[tag], v] if not isinstance(rs[tag], list) else (rs + [v])
                        else:
                            rs[tag] = v
            lst.append(rs)
    return lst


def get_body_by_flag(body, start_flag, end_flag=None):
    rs = []
    start, end = False, False
    for b in body:
        if isinstance(b, str) and b == '':
            continue
        if isinstance(b, str) and start_flag in b:
            start = True
        elif isinstance(b, str) and end_flag is not None and end_flag in b:
            end = True
        if start and not end:
            rs.append(b)
        elif end:
            return rs
    return rs


def get_single_body_by_flag(body, flag):
    find = False
    for b in body:
        if isinstance(b, str) and flag in b:
            find = True
            continue
        if find and isinstance(b, pd.DataFrame):
            return b
        elif find and isinstance(b, str) and b.strip() != '':
            return
    return


def read_identity(df):
    """
    身份信息
    -------------------------------
    0,1,2,3,4,5,6,7
    性别,出生日期,婚姻状况,手机号码,单位电话,住宅电话,学历,学位
    男性,1,已婚,,0222510,283291,"大学专科和专科学校（简称""大专""）",其他
    通讯地址,通讯地址,通讯地址,户籍地址,户籍地址,户籍地址,户籍地址,户籍地址
    天津市北辰,天津市北辰,天津市北辰,天津市天津市河,天津市天津市河,天津市天津市河,天津市天津市河,天津市天津市河
    -------------------------------
    :param df:
    :return:
    """
    if df is None:
        return
    if len(df.columns) != 8 or len(df.index) != 4:
        raise ValueError('身份信息数据有误或结构调整,请确认')
    identity = Identity()
    identity.gender = df.iloc[1, 0]
    identity.birthday = df.iloc[1, 1]
    identity.maritalState = df.iloc[1, 2]
    identity.mobile = df.iloc[1, 3]
    identity.officeTelephoneNo = df.iloc[1, 4]
    identity.homeTelephoneNo = df.iloc[1, 5]
    identity.eduLevel = df.iloc[1, 6]
    identity.eduDegree = df.iloc[1, 7]
    identity.postAddress = df.iloc[3, 0]
    identity.registeredAddress = df.iloc[3, 7]

    return identity


def read_residence(df):
    """
    0,1,2,3
    编号,居住地址,居住状况,信息更新日期
    1,中国天津市天津市（市辖区）05,其他,2017.09.29
    2,天津市北辰区佳,租房,2017.09.29
    3,天津市市辖区北辰区北辰区嘉阳0,按揭,2017.09.25
    4,小淀镇,未知,2017.09.18
    5,天津市北辰区小淀,按揭,2016.09.20
    :param df:
    :return:
    """
    if df is None:
        return
    if len(df.columns) != 4:
        raise ValueError("居住信息数据有误或结构调整,请确认")
    residence_lst = []
    mapping = {'编号': 'no', '居住地址': 'address', '居住状况': 'status', '信息更新日期': 'update_date'}
    df.columns = [mapping.get(v, v) for v in df.iloc[0]]
    df = df.drop([0], axis=0)
    df.set_index(['no'], inplace=True)
    for row in df.iterrows():
        r = Residence()
        r.getTime = row[1]['update_date']
        r.address = row[1]['address']
        r.residenceType = row[1]['status']
        residence_lst.append(r)
    return residence_lst


def read_spouse(df):
    """
    0,1,2,3,4
    姓名,证件类型,证件号码,工作单位,联系电话
    --,--,--,--,--
    :param df:
    :return:
    """
    if df is None:
        return None
    if len(df.columns) != 5 or len(df.index) != 2:
        raise ValueError('配偶信息有误或结构变化,请确认')
    spouse = Spouse()
    spouse.name = df.iloc[1, 0]
    spouse.certType = df.iloc[1, 1]
    spouse.certNo = df.iloc[1, 2]
    spouse.employer = df.iloc[1, 3]
    spouse.telephoneNo = df.iloc[1, 4]
    return spouse


def read_professional(body):
    """
    0,1,2,3,4,5,6
    编号,工作单位,工作单位,工作单位,工作单位,单位地址,单位地址
    1,天津市,天津市,天津市,天津市,--,--
    2,天津,天津,天津,天津,中国天津市天津,中国天津市天津
    3,天津市,天津市,天津市,天津市,天津市市辖区南开区天津市南开区鞍山西道学湖里8号,天津市市辖区南开区天津市南开区鞍山西道学湖里8号
    4,,,,,学湖里8号,学湖里8号
    5,,,,,天津市南开区广开五马路兴云里5号楼,天津市南开区广开五马路兴云里5号楼
    编号,职业,行业,职务,职称,进入本单位年份,信息更新日期
    1,办事人员和有关人员,--,中级领导（行政级别局级以下领导或大公司中级管理人员）,--,1985,2017.09.29
    2,商业、服务业人员,居民服务和其他服务业,一般员工,无,--,2017.09.29
    3,--,公共管理和社会组织,一般员工,--,1985,2017.09.25
    4,--,--,--,--,2011,2017.09.18
    5,--,--,其他,--,--,2016.09.20
    :param body:
    :return:
    """
    professional_lst = []
    cls = 0
    counter = 0
    if body is None:
        return professional_lst
    for row in body.iterrows():
        if '工作单位' in list(row[1]):
            cls = 1
            continue
        elif '职业' in list(row[1]):
            cls = 2
            continue
        if cls == 1:
            professional = Professional()
            professional.employer = row[1][1]
            professional.employerAddress = row[1][6]
        elif cls == 2:
            professional = professional_lst[counter]
            professional.occupation = row[1][1]
            professional.industry = row[1][2]
            professional.duty = row[1][3]
            professional.title = row[1][4]
            professional.startYear = row[1][5]
            professional.getTime = row[1][6]
            counter += 1
        professional_lst.append(professional)
    return professional_lst


def read_personal_basic_info(body):
    """
    :param body:
    :return:
    """
    personal = PersonalBasicInfo()
    personal_body = get_body_by_flag(body, '个人基本信息', '信息概要')
    identity_body = get_single_body_by_flag(personal_body, '身份信息')
    professional_body = get_single_body_by_flag(personal_body, '职业信息')
    residence_body = get_single_body_by_flag(personal_body, '居住信息')
    spouse_body = get_single_body_by_flag(personal_body, '配偶信息')
    personal.identity = read_identity(identity_body)
    personal.residence = read_residence(residence_body)
    personal.spouse = read_spouse(spouse_body)
    personal.professional = read_professional(professional_body)

    return personal


def read_credit_cue(body):
    """
    0,1,2,3,4,5,6,7,8,9
    个人住房贷款笔数,个人商用房（包括商住两用）贷款笔数,其他贷款笔数,首笔贷款发放月份,贷记卡账户数,首张贷记卡发卡月份,准贷记卡账户数,首张准贷记卡发卡月份,本人声明数目,异议标注数目
    1,0,17,2012.11,18,2007.01,2,2012.03,0,0
    :param body:
    :return:
    """
    if body is None:
        return
    credit_cue = CreditCue()
    credit_cue.perHouseLoanCount = body.iloc[1, 0]
    credit_cue.perBusinessHouseLoanCount = body.iloc[1, 1]
    credit_cue.otherLoanCount = body.iloc[1, 2]
    credit_cue.firstLoanOpenMonth = body.iloc[1, 3]
    credit_cue.loanCardCount = body.iloc[1, 4]
    credit_cue.firstLoanCardOpenMonth = body.iloc[1, 5]
    credit_cue.standardLoanCardCount = body.iloc[1, 6]
    credit_cue.firstStandardLoanCardOpenMonth = body.iloc[1, 7]
    credit_cue.announceCount = body.iloc[1, 8]
    credit_cue.dissentCount = body.iloc[1, 9]
    return credit_cue


def read_overdue_and_fell_back(body):
    """
    ------------------------
    0,1,2,3,4,5,6,7
    呆账信息汇总,呆账信息汇总,资产处置信息汇总,资产处置信息汇总,保证人代偿信息汇总,保证人代偿信息汇总,保证人代偿信息汇总,保证人代偿信息汇总
    笔数,余额,笔数,余额,笔数,余额,,
    1,"148,500",0,0,3,"521,405",,
    ------------------------
    0,1,2,3,4,5,6,7,8,9,10,11
    贷款逾期,贷款逾期,贷款逾期,贷款逾期,贷记卡逾期,贷记卡逾期,贷记卡逾期,贷记卡逾期,准贷记卡60天以上透支,准贷记卡60天以上透支,准贷记卡60天以上透支,准贷记卡60天以上透支
    "笔数","月份数","单月最高逾期总额","最长逾期月数","账户数","月份数","单月最高逾期总额","最长逾期月数","账户数","月份数","单月最高透支余额","最长透支月数"
    7,10,"36,195",7,6,9,"309,422",7,0,0,0,0
    :param body:
    :return:
    """
    if body is None:
        return
    overdue_and_fellback = OverdueAndFellBack()
    # 违约信息概要
    fell_back_body = get_single_body_by_flag(body, '逾期及违约信息概要')
    if fell_back_body is not None:
        fell_back_summary = FellBackSummary()
        fell_back_summary.fellBackDebtSumCount = fell_back_body.iloc[2, 0]
        fell_back_summary.fellBackDebtSumBalance = fell_back_body.iloc[2, 1]
        fell_back_summary.assetDispositionSumCount = fell_back_body.iloc[2, 2]
        fell_back_summary.assetDispositionSumBalance = fell_back_body.iloc[2, 3]
        fell_back_summary.assureerRepaySumCount = fell_back_body.iloc[2, 4]
        fell_back_summary.assureerRepaySumBalance = fell_back_body.iloc[2, 5]
        overdue_and_fellback.fellBackSummary = fell_back_summary

    # 逾期信息概要
    overdue_body = get_single_body_by_flag(body, '逾期（透支）信息汇总')
    if overdue_body is not None:
        overdue = OverdueSummary()
        overdue.loanSumCount = overdue_body.iloc[2, 0]
        overdue.loanSumMonths = overdue_body.iloc[2, 1]
        overdue.loanSumHighestOverdueAmountPerMon = overdue_body.iloc[2, 2]
        overdue.loanSumMaxDuration = overdue_body.iloc[2, 3]
        overdue.loanCardSumCount = overdue_body.iloc[2, 4]
        overdue.loanCardSumMonths = overdue_body.iloc[2, 5]
        overdue.loanCardSumHighestOverdueAmountPerMon = overdue_body.iloc[2, 6]
        overdue.loanCardSumMaxDuration = overdue_body.iloc[2, 7]
        overdue.standardLoanCardSumCount = overdue_body.iloc[2, 8]
        overdue.standardLoanCardSumMonths = overdue_body.iloc[2, 9]
        overdue.standardLoanCardSumHighestOverdueAmountPerMon = overdue_body.iloc[2, 10]
        overdue.standardLoanCardSumMaxDuration = overdue_body.iloc[2, 11]
        overdue_and_fellback.overdueSummary = overdue

    return overdue_and_fellback


def read_share_and_debt(body):
    """
    ------------------------
    0,1,2,3,4,5
    贷款法人机构数,贷款机构数,笔数,合同总额,余额,最近6个月平均应还款
    8,8,12,"1,503,400","1,206,316","26,209"
    ------------------------
    0,1,2,3,4,5,6,7
    "发卡法人机构数","发卡机构数",账户数,授信总额,"单家行最高授信额","单家行最低授信额",已用额度,"最近6个月平均使用额度"
    9,9,13,"709,172","233,377","14,000","520,364","487,711"
    ------------------------
    0,1,2,3,4,5,6,7
    "发卡法人机构数","发卡机构数",账户数,授信总额,"单家行最高授信额","单家行最低授信额",透支余额,"最近6个月平均透支余额"
    1,1,2,0,"50,000","50,000",0,0

    :param body:
    :return:
    """
    share_and_debt = ShareAndDebt()
    # 未结清贷款信息汇总
    un_paid_loan_body = get_single_body_by_flag(body, '未结清贷款信息汇总')
    if un_paid_loan_body is not None:
        un_paid_loan = ShareAndDebtCommon()
        un_paid_loan.financeCorpCount = un_paid_loan_body.iloc[1, 0]
        un_paid_loan.financeOrgCount = un_paid_loan_body.iloc[1, 1]
        un_paid_loan.accountCount = un_paid_loan_body.iloc[1, 2]
        un_paid_loan.creditLimit = un_paid_loan_body.iloc[1, 3]
        un_paid_loan.balance = un_paid_loan_body.iloc[1, 4]
        un_paid_loan.latest6MonthUsedAvgAmount = un_paid_loan_body.iloc[1, 5]
        share_and_debt.unPaidLoan = un_paid_loan

    # 未销户贷记卡信息汇总
    un_destroy_loan_card_body = get_single_body_by_flag(body, '未销户贷记卡信息汇总')
    if un_destroy_loan_card_body is not None:
        un_destroy_loan_card = ShareAndDebtCommon()
        un_destroy_loan_card.financeCorpCount = un_destroy_loan_card_body.iloc[1, 0]
        un_destroy_loan_card.financeOrgCount = un_destroy_loan_card_body.iloc[1, 1]
        un_destroy_loan_card.accountCount = un_destroy_loan_card_body.iloc[1, 2]
        un_destroy_loan_card.creditLimit = un_destroy_loan_card_body.iloc[1, 3]
        un_destroy_loan_card.maxCreditLimitPerOrg = un_destroy_loan_card_body.iloc[1, 4]
        un_destroy_loan_card.minCreditLimitPerOrg = un_destroy_loan_card_body.iloc[1, 5]
        un_destroy_loan_card.usedCreditLimit = un_destroy_loan_card_body.iloc[1, 6]
        un_destroy_loan_card.latest6MonthUsedAvgAmount = un_destroy_loan_card_body.iloc[1, 7]
        share_and_debt.unDestroyLoanCard = un_destroy_loan_card

    # 未销户准贷记卡信息汇总
    un_destroy_standard_loan_card_body = get_single_body_by_flag(body, '未销户准贷记卡信息汇总')
    if un_destroy_standard_loan_card_body is not None:
        un_destroy_standard_loan_card = ShareAndDebtCommon()
        un_destroy_standard_loan_card.financeCorpCount = un_destroy_standard_loan_card_body.iloc[1, 0]
        un_destroy_standard_loan_card.financeOrgCount = un_destroy_standard_loan_card_body.iloc[1, 1]
        un_destroy_standard_loan_card.accountCount = un_destroy_standard_loan_card_body.iloc[1, 2]
        un_destroy_standard_loan_card.creditLimit = un_destroy_standard_loan_card_body.iloc[1, 3]
        un_destroy_standard_loan_card.maxCreditLimitPerOrg = un_destroy_standard_loan_card_body.iloc[1, 4]
        un_destroy_standard_loan_card.minCreditLimitPerOrg = un_destroy_standard_loan_card_body.iloc[1, 5]
        un_destroy_standard_loan_card.usedCreditLimit = un_destroy_standard_loan_card_body.iloc[1, 6]
        un_destroy_standard_loan_card.latest6MonthUsedAvgAmount = un_destroy_standard_loan_card_body.iloc[1, 7]
        share_and_debt.unDestroyStandardLoanCard = un_destroy_standard_loan_card
    return share_and_debt


def read_summary_info(body):
    summary_info = SummaryInfo()
    summary_info_body = get_body_by_flag(body, '信息概要', '信贷交易信息明细')
    credit_cue_body = get_single_body_by_flag(summary_info_body, '信用提示')

    summary_info.creditCue = read_credit_cue(credit_cue_body)
    summary_info.overdueAndFellBack = read_overdue_and_fell_back(summary_info_body)
    summary_info.shareAndDebt = read_share_and_debt(summary_info_body)

    return summary_info


def read_assurer_repay(body):
    """
    0,1,2,3,4,5
    编号,代偿机构,最近一次代偿日期,累计代偿金额,最近一次还款日期,余额
    1,YS,2018.01.29,"433,883",2018.01.30,"423,527"
    2,DT,2018.04.11,"71,886",--,"71,886"
    3,FB,2018.05.24,"25,992",--,"25,992"
    :param body:
    :return:
    """
    rs = []
    assurer_repay_body = get_single_body_by_flag(body, '）保证人代偿信息')
    if assurer_repay_body is None:
        return {}
    for row in assurer_repay_body.iterrows():
        row = row[1]
        if '代偿机构' in list(row):
            continue
        else:
            assurerRepay = AssurerRepay()
            assurerRepay.org = row[1]
            assurerRepay.accumulativeAssurerRepayAmount = row[2]
            assurerRepay.recentAssurerRepayDate = row[3]
            assurerRepay.recentRepayDate = row[4]
            assurerRepay.balance = row[5]
            rs.append(assurerRepay)
    return rs


def read_guarantee_info(body):
    guarantee_info = GuaranteeInfo()
    guarantee_info.guarantee = []
    guarantee_info_body = get_single_body_by_flag(body, '对外贷款担保信息')
    if guarantee_info_body is not None:
        guarantee_info.guaranteeFormat = '对外贷款担保信息'
        for row in guarantee_info_body.iterrows():
            row = row[1]
            if '编号' in row[0]:
                continue
            guarantee = Guarantee()
            guarantee.organname = row[1]
            guarantee.contractMoney = row[2]
            guarantee.beginDate = row[3]
            guarantee.endDate = row[4]
            guarantee.guananteeMoney = row[5]
            guarantee.guaranteeBalance = row[6]
            guarantee.class5State = row[7]
            guarantee.billingDate = row[7]
            guarantee_info.guarantee.append(guarantee)
        return guarantee_info
    else:
        return {}


def read_loan(body):
    """
    账户状态	五级分类	本金余额	剩余还款期数	本月应还款 	应还款日 	本月实还款 	最近一次还款日期
    正常	正常	445,607	213	3,738	2019.09.30	3,738	2019.09.20
    当前逾期期数	当前逾期金额	逾期31-60天未还本金	逾期61－90天未还本金	逾期91－180天未还本金	逾期180天以上未还本金
    0	0	0	0	0	0
    2017年10月-2019年09月的还款记录
    N	N	N	N	N	N	N	N	N	N	N	N	N	N	N	N	N	N	N	N	N	N	N	N

    2014年11月-2016年12月的逾期记录
    逾期月份	逾期持续月数	逾期金额	逾期月份	逾期持续月数	逾期金额
    2016.04	7	430,003	2016.03	7	422,472
    2016.02	7	415,543	2016.01	7	408,266
    2015.12	7	401,116	2015.11	7	394,314

    特殊交易类型	发生日期	变更月数	发生金额	明细记录
    其他	2013.05.29	0	295,516	提前还款
    :param body:
    :return:
    """
    loan_body_lst = get_body_by_flag(body, '贷款', '贷记卡')
    loan_lst = []
    find = False
    counter = -1
    for ln in loan_body_lst:
        if isinstance(ln, str) and ('（二）' in ln or '（一）' in ln or '（三）' in ln or '（四）' in ln):
            continue
        if isinstance(ln, str) and '贷款' in ln:
            counter += 1
            find = True
            loan = Loan()
            loan.statements = ln
            loan_lst.append(loan)
        elif isinstance(ln, pd.DataFrame) and find:
            find = False
            if len(ln.index) == 0:
                continue
            if ln.shape == (1, 1) and ln.iloc[0, 0] == '':
                continue
            cls = None
            row_title = None
            for ii, row in enumerate(ln.iterrows()):
                row = row[1]
                if '账户' in row[0] and '状态' in row[0]:
                    row_title = row
                    cls = 1
                    continue
                elif '当前逾期期数' in row[0]:
                    row_title = row
                    cls = 2
                    continue
                elif '逾期记录' in row[0] or '逾期月份' in row[0]:
                    row_title = row
                    cls = 3
                    continue
                elif '特殊交易类型' in row[0]:
                    row_title = row
                    cls = 4
                    continue
                elif '还款记录' in row[0]:
                    row_title = row
                    cls = 5
                    latest24Date = row[0]
                    continue
                if cls == 1:
                    loan = loan_lst[counter]
                    loan.state = row[0]
                    loan.class5State = row[3]
                    loan.balance = row[6]
                    loan.remainPaymentCyc = row[9]
                    loan.scheduledPaymentAmount = row[12]
                    loan.scheduledPaymentDate = row[15]
                    loan.actualPaymentAmount = row[18]
                    loan.recentPayDate = row[21]
                elif cls == 2:
                    loan = loan_lst[counter]
                    loan.currOverdueCyc = row[0]
                    loan.currOverdueAmount = row[4]
                    loan.overdue31To60Amount = row[8]
                    loan.overdue61To90Amount = row[12]
                    loan.overdue91To180Amount = row[16]
                    loan.overdueOver180Amount = row[20]
                elif cls == 5:
                    loan = loan_lst[counter]
                    loan.latest24Date = latest24Date
                    loan.latest24State = ''.join(row)
                elif cls == 3:
                    loan = loan_lst[counter]
                    gap = 4 if len(row) == 24 else 1
                    detail = OverdueRecordDetail()
                    detail.month = row[0 * gap]
                    detail.lastMonths = row[1 * gap]
                    detail.amount = row[2 * gap]
                    if loan.overdueRecord is not None and isinstance(loan.overdueRecord.overdueRecordDetail, list):
                        loan.overdueRecord.overdueRecordDetail.append(detail)
                    else:
                        loan.overdueRecord = OverdueRecord()
                        loan.overdueRecord.overdueRecordDetail = [detail]
                    detail = OverdueRecordDetail()
                    for jj, tl in enumerate(row_title):
                        if jj < 3:
                            continue
                        if '逾期月份' in tl:
                            detail.month = row[jj]
                        elif '逾期持续月数' in tl:
                            detail.lastMonths = row[jj]
                        elif '逾期金额' in tl:
                            detail.amount = row[jj]
                    # detail.month = row[3 * gap]
                    # detail.lastMonths = row[4 * gap]
                    # detail.amount = row[5 * gap]
                    loan.overdueRecord.overdueRecordDetail.append(detail)
                elif cls == 4:
                    loan = loan_lst[counter]
                    gap = 4 if len(row) == 24 else 1
                    special = SpecialRecord()
                    special.tradeType = row[0 * gap]
                    special.date = row[1 * gap]
                    special.changeMonths = row[2 * gap]
                    special.amount = row[3 * gap]
                    special.detail = row[4 * gap]
                    if loan.specials is not None:
                        loan.specials.append(special)
                    else:
                        loan.specials = [special]
    return loan_lst


def read_loan_card(loan_card_body_lst):
    """
    账户状态	已用额度	最近6个月平均使用额度	最大使用额度	本月应还款
    正常	31,371	16,809 	86,308	402
    账单日 	本月实还款 	最近一次还款日期	当前逾期期数	当前逾期金额
    2019.09.14	4,020	2019.09.07	0	0
    2017年10月-2019年09月的还款记录
    N	N	*	N	*	*	*	*	*	N	N	*	*	*	*	*	*	*	*	*	N	N	N	N

    2014年11月-2018年05月的逾期记录
    逾期月份	逾期持续月数	逾期金额	逾期月份	逾期持续月数	逾期金额
    2017.02	3	56	2017.01	2	35
    2016.12	1	15	--	--	--

    :param loan_card_body_lst:
    :return:
    """
    loan_card_lst = []
    find = False
    counter = -1
    for lc in loan_card_body_lst:
        if isinstance(lc, str) and ('一）' in lc or '二）' in lc or '三）' in lc or '四）' in lc):
            continue
        if isinstance(lc, str) and '贷记卡' in lc and not find:
            find = True
            counter += 1
            loan = LoanCard()
            loan.statements = lc
            loan_card_lst.append(loan)
        elif isinstance(lc, pd.DataFrame) and find:
            find = False
            if len(lc.index) == 0:
                continue
            if lc.shape == (1, 1) and lc.iloc[0, 0] == '':
                continue
                row_title = None
            for row in lc.iterrows():
                row = row[1]
                if '账户' in row[0] and '状态' in row[0]:
                    row_title = row
                    cls = 1
                    continue
                elif '账单日' in row[0]:
                    row_title = row
                    cls = 2
                    continue
                elif '逾期记录' in row[0] or '逾期月份' in row[0]:
                    row_title = row
                    cls = 3
                    continue
                elif '特殊交易类型' in row[0]:
                    row_title = row
                    cls = 4
                    continue
                elif '还款记录' in row[0]:
                    row_title = row
                    cls = 5
                    latest24Date = row[0]
                    continue

                if cls == 1:
                    loan = loan_card_lst[counter]
                    loan.state = row[0]
                    loan.usedCreditLimitAmount = row[4]
                    loan.latest6MonthUsedAvgAmount = row[8]
                    loan.usedHighestAmount = row[16]
                    loan.scheduledPaymentAmount = row[20]
                elif cls == 2:
                    loan = loan_card_lst[counter]
                    loan.scheduledPaymentDate = row[0]
                    loan.actualPaymentAmount = row[4]
                    loan.recentPayDate = row[8]
                    loan.currOverdueCyc = row[16]
                    loan.currOverdueAmount = row[20]
                elif cls == 5:
                    loan = loan_card_lst[counter]
                    loan.latest24Date = latest24Date
                    loan.latest24State = ''.join(row)
                elif cls == 3:
                    loan = loan_card_lst[counter]
                    gap = 4 if len(lc.columns) == 24 else 1
                    detail = OverdueRecordDetail()
                    detail.month = row[0 * gap]
                    detail.lastMonths = row[1 * gap]
                    detail.amount = row[2 * gap]
                    if loan.overdueRecord is not None and isinstance(loan.overdueRecord.overdueRecordDetail, list):
                        loan.overdueRecord.overdueRecordDetail.append(detail)
                    else:
                        loan.overdueRecord = OverdueRecord()
                        loan.overdueRecord.overdueRecordDetail = [detail]
                    detail = OverdueRecordDetail()
                    for jj, tl in enumerate(row_title):
                        if jj < 3:
                            continue
                        if '逾期月份' in tl:
                            detail.month = row[jj]
                        elif '逾期持续月数' in tl:
                            detail.lastMonths = row[jj]
                        elif '逾期金额' in tl:
                            detail.amount = row[jj]
                    loan.overdueRecord.overdueRecordDetail.append(detail)
                elif cls == 4:
                    loan = loan_card_lst[counter]
                    gap = 4 if len(row) == 24 else 1
                    special = SpecialRecord()
                    special.tradeType = row[0 * gap]
                    special.date = row[1 * gap]
                    special.changeMonths = row[2 * gap]
                    special.amount = row[3 * gap]
                    special.detail = row[4 * gap]
                    if loan.specials is not None:
                        loan.specials.append(special)
                    else:
                        loan.specials = [special]
    return loan_card_lst


def read_standard_loan_card(loan_card_body_lst):
    loan_card_lst = []
    find = False
    counter = -1
    cls = None
    for lc in loan_card_body_lst:
        if isinstance(lc, str) and ('一）' in lc or '二）' in lc or '三）' in lc or '四）' in lc):
            continue
        if isinstance(lc, str) and '准贷记卡' in lc and not find:
            find = True
            counter += 1
            loan = LoanCard()
            loan.statements = lc
            loan_card_lst.append(loan)
        elif isinstance(lc, pd.DataFrame) and find:
            find = False
            if len(lc.index) == 0:
                continue
            if lc.shape == (1, 1) and lc.iloc[0, 0] == '':
                continue
            for ii, row in enumerate(lc.iterrows()):
                row = row[1]
                if '账户' in row[0] and '状态' in row[0]:
                    cls = 1
                    continue
                elif '逾期记录' in row[0] or '逾期月份' in row[0]:
                    cls = 3
                    continue
                elif '特殊交易类型' in row[0]:
                    cls = 4
                    continue
                elif '还款记录' in row[0]:
                    cls = 5
                    latest24Date = row[0]
                    continue
                loan = loan_card_lst[counter]
                if cls == 1:
                    loan.state = row[0]
                    loan.usedCreditLimitAmount = row[2]
                    loan.latest6MonthUsedAvgAmount = row[4]
                    loan.usedHighestAmount = row[8]
                    loan.scheduledPaymentDate = row[11]
                    loan.actualPaymentAmount = row[14]
                    loan.recentPayDate = row[17]
                    loan.due180pAmount = row[20]
                elif cls == 5:
                    loan.latest24Date = latest24Date
                    loan.latest24State = ''.join(row)
                elif cls == 3:
                    gap = 4 if len(lc.columns) == 24 else 1
                    detail = OverdueRecordDetail()
                    detail.month = lc.iloc[ii, 0 * gap]
                    detail.lastMonths = lc.iloc[ii, 1 * gap]
                    detail.amount = lc.iloc[ii, 2 * gap]
                    if loan.overdueRecord is not None and isinstance(loan.overdueRecord.overdueRecordDetail, list):
                        loan.overdueRecord.overdueRecordDetail.append(detail)
                    else:
                        loan.overdueRecord = OverdueRecord()
                        loan.overdueRecord.overdueRecordDetail = [detail]
                    detail = OverdueRecordDetail()
                    detail.month = lc.iloc[ii, 3 * gap]
                    detail.lastMonths = lc.iloc[ii, 4 * gap]
                    detail.amount = lc.iloc[ii, 5 * gap]
                    loan.overdueRecord.overdueRecordDetail.append(detail)
                elif cls == 4:
                    gap = 4 if len(row) == 24 else 1
                    special = SpecialRecord()
                    special.tradeType = row[0 * gap]
                    special.date = row[1 * gap]
                    special.changeMonths = row[2 * gap]
                    special.amount = row[3 * gap]
                    special.detail = row[4 * gap]
                    if loan.specials is not None:
                        loan.specials.append(special)
                    else:
                        loan.specials = [special]
    return loan_card_lst


def read_credit_detail(body):
    """
    信贷交易信息明细
    :return:
    """
    credit_detail = CreditDetail()
    credit_detail_body = get_body_by_flag(body, '信贷交易信息明细', '公共信息明细')

    # 保证人代偿信息
    credit_detail.assurerRepay = read_assurer_repay(credit_detail_body)
    # 担保信息
    credit_detail.guaranteeInfo = read_guarantee_info(credit_detail_body)

    # 贷款
    credit_detail.loan = read_loan(credit_detail_body)
    # 贷记卡
    loan_card_body_lst = get_body_by_flag(credit_detail_body, '）贷记卡', '）准贷记卡')
    credit_detail.loanCard = read_loan_card(loan_card_body_lst)
    # 准贷记卡
    standard_loan_card_body_lst = get_body_by_flag(credit_detail_body, '）准贷记卡')
    credit_detail.standardLoanCard = read_standard_loan_card(standard_loan_card_body_lst)
    return credit_detail


def read_acc_fund(body):
    acc_fund_lst = []
    acc_fund_body = get_single_body_by_flag(body, '住房公积金参缴记录')
    counter = 0
    if acc_fund_body is not None:
        for row in acc_fund_body.iterrows():
            row = row[1]
            if '参缴地' in row[1]:
                cls = 1
                continue
            elif '缴费单位' in row[1]:
                cls = 2
                continue
            if cls == 1:
                acc_fund = AccFund()
                acc_fund.area = row[1]
                acc_fund.registerDate = row[2]
                acc_fund.firstMonth = row[3]
                acc_fund.toMonth = row[4]
                acc_fund.state = row[5]
                acc_fund.pay = row[6]
                acc_fund.ownPercent = row[7]
                acc_fund.comPercent = row[8]
                acc_fund_lst.append(acc_fund)
            elif cls == 2:
                acc_fund = acc_fund_lst[counter]
                acc_fund.organname = row[1]
                acc_fund.getTime = row[9]
                counter += 1
    return acc_fund_lst


def read_public_info(body):
    public_info = PublicInfo()
    public_info_body = get_body_by_flag(body, '四 公共信息明细', '五 查询记录')
    public_info.accFund = read_acc_fund(public_info_body)

    return public_info


def read_query_record_detail(body):
    rs = []
    record_detail_body = get_single_body_by_flag(body, '信贷审批查询记录明细')
    find = False
    for row in record_detail_body.iterrows():
        row = row[1]
        if '查询日期' in list(row):
            find = True
            continue
        if find:
            queryRecordDetail = QueryRecordDetail()
            queryRecordDetail.queryDate = row[1]
            queryRecordDetail.querier = row[2]
            queryRecordDetail.queryReason = row[3]
            rs.append(queryRecordDetail)
    return rs


def read_query_record_summary(body):
    record_summary_body = get_single_body_by_flag(body, '查询记录汇总')
    record_summary = QueryRecordSummary()
    record_summary.latestMonthQueryorgSumLoanApproval = record_summary_body.iloc[2, 0]
    record_summary.latestMonthQueryorgSumLoanCardApproval = record_summary_body.iloc[2, 1]
    record_summary.latestMonthQueryRecordSumLoanApproval = record_summary_body.iloc[2, 2]
    record_summary.latestMonthQueryRecordSumLoanCardApproval = record_summary_body.iloc[2, 3]
    record_summary.latestMonthQueryRecordSumPersonal = record_summary_body.iloc[2, 4]
    record_summary.twoYearQueryRecordSumCollection = record_summary_body.iloc[2, 5]
    record_summary.twoYearQueryRecordSumGuarantee = record_summary_body.iloc[2, 6]
    record_summary.twoYearQueryRecordSumSpecial = record_summary_body.iloc[2, 7]

    return record_summary


def read_query_record(body):
    query_record = QueryRecord()
    query_record_body = get_body_by_flag(body, '五 查询记录', '报告说明')
    # 查询记录汇总
    query_record.recordSummary = read_query_record_summary(query_record_body)

    # 信贷审批查询记录明细
    query_record.recordInfo = read_query_record_detail(query_record_body)
    return query_record


def to_json(word_file, json_file):
    document = docx.Document(word_file)
    body = prefix_word(document)
    obj = PBOCEntity()
    obj.body_str = str(body).replace('\n', '')  # 提供文字版报告
    # 报告基本信息
    obj.header = read_report_info(body)
    # 个人基本信息
    obj.personalInfo = read_personal_basic_info(body)
    # 信息概要
    obj.summary_info = read_summary_info(body)
    # 信贷交易信息明细
    obj.creditDetail = read_credit_detail(body)
    # 公共信息
    obj.publicInfo = read_public_info(body)
    # 查询记录
    obj.queryRecord = read_query_record(body)

    with open(json_file, 'w', encoding='utf-8') as of:
        obj = obj_to_dict(obj)
        json.dump(obj, of, ensure_ascii=False)


def logger_(log_file_name=None, name=__name__, stdout_on=True):
    log_file_name = log_file_name if log_file_name is not None else (__name__ + '.log')

    logger_ = logging.getLogger(name)
    fmt = '%(msg)s'
    date_fmt = '%Y-%m-%d %H:%M:%S'
    formatter = logging.Formatter(fmt=fmt, datefmt=date_fmt)

    if stdout_on:
        stout_handler = logging.StreamHandler(sys.stdout)
        stout_handler.setFormatter(formatter)
        logger_.addHandler(stout_handler)

    file_handler = logging.FileHandler(log_file_name, encoding='utf-8')
    file_handler.setFormatter(formatter)
    logger_.addHandler(file_handler)

    logger_.setLevel(logging.DEBUG)

    return logger_


def read_file_names(file_path):
    table_names = []
    if os.path.exists(file_path):
        file_names = os.listdir(file_path)
        for file_i in file_names:
            path_file = os.path.join(file_path, file_i)
            table_names.append(path_file)
        if len(table_names) == 0:
            print('there is no file in the path')
    else:
        print('the path [{}] is not exist!'.format(file_path))
    return table_names


if __name__ == '__main__':
    now_str = datetime.now().strftime('%Y%m%d%H%M%S')

    log_dir = os.path.join(os.path.abspath(os.getcwd()), 'log', datetime.now().strftime('%Y%m%d'))
    os.makedirs(log_dir, exist_ok=True)
    try:
        args = docopt(__doc__)
        print(args)
        is_extract = args['extract']

        log_file = os.path.join(log_dir, 'word2json_{0}.log'.format(now_str))
        logger = logger_(log_file, name='tojson')
        word_file = args['<word_file>']
        json_file = args['<json_file>']
        if json_file is None:
            json_file = '{0}.json'.format(word_file)
        to_json(word_file, json_file)

    except Exception as e:
        logger.error(traceback.format_exc())
